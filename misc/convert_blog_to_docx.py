#!/usr/bin/env python3
"""Convert the blog article markdown to a formatted Word document."""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, font_size=None, color=None, space_after=None, space_before=None, alignment=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if alignment is not None:
        p.alignment = alignment
    return p

def add_rich_paragraph(doc, parts, style='Normal', space_after=None, space_before=None, alignment=None):
    """Add a paragraph with mixed formatting (bold, italic, normal parts)."""
    p = doc.add_paragraph(style=style)
    for part in parts:
        text = part.get('text', '')
        run = p.add_run(text)
        run.bold = part.get('bold', False)
        run.italic = part.get('italic', False)
        if 'font_size' in part:
            run.font.size = Pt(part['font_size'])
        if 'color' in part:
            run.font.color.rgb = RGBColor(*part['color'])
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if alignment is not None:
        p.alignment = alignment
    return p

def parse_inline(text):
    """Parse inline markdown (bold, italic, code, links) into parts."""
    parts = []
    i = 0
    while i < len(text):
        # Bold + italic
        if text[i:i+3] == '***':
            end = text.find('***', i+3)
            if end != -1:
                parts.append({'text': text[i+3:end], 'bold': True, 'italic': True})
                i = end + 3
                continue
        # Bold
        if text[i:i+2] == '**':
            end = text.find('**', i+2)
            if end != -1:
                parts.append({'text': text[i+2:end], 'bold': True})
                i = end + 2
                continue
        # Italic
        if text[i] == '*' and (i == 0 or text[i-1] != '*') and (i+1 < len(text) and text[i+1] != '*'):
            end = text.find('*', i+1)
            if end != -1 and (end+1 >= len(text) or text[end+1] != '*'):
                parts.append({'text': text[i+1:end], 'italic': True})
                i = end + 1
                continue
        # Inline code
        if text[i] == '`':
            end = text.find('`', i+1)
            if end != -1:
                parts.append({'text': text[i+1:end], 'bold': True})
                i = end + 1
                continue
        # Link [text](url) - just use the text
        if text[i] == '[':
            bracket_end = text.find(']', i+1)
            if bracket_end != -1 and bracket_end+1 < len(text) and text[bracket_end+1] == '(':
                paren_end = text.find(')', bracket_end+2)
                if paren_end != -1:
                    link_text = text[i+1:bracket_end]
                    parts.append({'text': link_text, 'italic': True})
                    i = paren_end + 1
                    continue
        # Regular character - accumulate
        if parts and not parts[-1].get('bold') and not parts[-1].get('italic'):
            parts[-1]['text'] += text[i]
        else:
            parts.append({'text': text[i]})
        i += 1
    return parts

def create_blog_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Modify styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hstyle = doc.styles[f'Heading {level}']
        hstyle.font.name = 'Calibri'
        hstyle.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
        if level == 1:
            hstyle.font.size = Pt(24)
            hstyle.paragraph_format.space_before = Pt(24)
            hstyle.paragraph_format.space_after = Pt(12)
        elif level == 2:
            hstyle.font.size = Pt(18)
            hstyle.paragraph_format.space_before = Pt(20)
            hstyle.paragraph_format.space_after = Pt(10)
        else:
            hstyle.font.size = Pt(14)
            hstyle.paragraph_format.space_before = Pt(16)
            hstyle.paragraph_format.space_after = Pt(8)

    # ── Title ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('Your Enterprise Data Is Already AI-Ready')
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x0a, 0x3d, 0x6b)
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run('Azure NetApp Files and the Agentic AI Revolution')
    sub_run.bold = True
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = RGBColor(0x1a, 0x56, 0x8e)
    subtitle_p.paragraph_format.space_after = Pt(12)

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline_p.add_run(
        'How a focused set of S3-compatible operations unlocks every major agentic AI framework\n'
        '— from RAG pipelines to MCP servers to autonomous multi-agent systems'
    )
    tag_run.italic = True
    tag_run.font.size = Pt(11)
    tag_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    tagline_p.paragraph_format.space_after = Pt(24)

    # ── Horizontal rule ──
    doc.add_paragraph('─' * 80).paragraph_format.space_after = Pt(12)

    # ── Read the markdown and parse ──
    with open('/home/user/ANF-OneLake-AIFoundry/misc/blog-anf-object-api-agentic-ai.md', 'r') as f:
        md = f.read()

    lines = md.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    skip_title = True  # skip the first H1 and subtitle

    while i < len(lines):
        line = lines[i]

        # Skip the title block (first H1, subtitle, first ---)
        if skip_title:
            if line.startswith('# ') or line.startswith('*How a focused') or line.strip() == '---':
                i += 1
                if line.strip() == '---':
                    skip_title = False
                continue

        # Code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                # Add code as a formatted block
                code_text = '\n'.join(code_lines)
                code_p = doc.add_paragraph()
                code_p.paragraph_format.space_before = Pt(8)
                code_p.paragraph_format.space_after = Pt(8)
                code_p.paragraph_format.left_indent = Cm(1)
                code_run = code_p.add_run(code_text)
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Parse table row
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # Skip separator rows
            if cells and all(set(c.strip()).issubset(set('-: ')) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table - render it
            in_table = False
            if table_rows:
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clean markdown from cell text
                            clean = cell_text.replace('**', '').replace('`', '').replace('*', '')
                            cell.text = clean
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(2)
                                paragraph.paragraph_format.space_before = Pt(2)
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                    run.font.name = 'Calibri'
                            # Bold header row
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                set_cell_shading(cell, "1a568e")
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                doc.add_paragraph()  # spacing after table
                table_rows = []
            # Don't increment i, process current line normally

        # Horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('─' * 80).paragraph_format.space_after = Pt(8)
            i += 1
            continue

        # Headings
        if line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue
        if line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        if line.startswith('# '):
            heading_text = line[2:].strip()
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+\*\*(.+?)\*\*\s*[—–-]\s*(.*)', line)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            run_bold = p.add_run(num_match.group(2))
            run_bold.bold = True
            run_normal = p.add_run(f' — {num_match.group(3)}')
            i += 1
            continue

        num_match2 = re.match(r'^(\d+)\.\s+\*\*(.+?)\*\*(.*)', line)
        if num_match2:
            p = doc.add_paragraph(style='List Number')
            run_bold = p.add_run(num_match2.group(2))
            run_bold.bold = True
            rest = num_match2.group(3)
            if rest:
                parts = parse_inline(rest)
                for part in parts:
                    run = p.add_run(part['text'])
                    run.bold = part.get('bold', False)
                    run.italic = part.get('italic', False)
            i += 1
            continue

        simple_num = re.match(r'^(\d+)\.\s+(.*)', line)
        if simple_num:
            p = doc.add_paragraph(style='List Number')
            parts = parse_inline(simple_num.group(2))
            for part in parts:
                run = p.add_run(part['text'])
                run.bold = part.get('bold', False)
                run.italic = part.get('italic', False)
            i += 1
            continue

        # Bullet lists
        bullet_match = re.match(r'^[-*]\s+\*\*(.+?)\*\*[:\s]*(.*)', line)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            run_bold = p.add_run(bullet_match.group(1))
            run_bold.bold = True
            rest = bullet_match.group(2)
            if rest:
                run_rest = p.add_run(f': {rest}' if not rest.startswith(':') else rest)
            i += 1
            continue

        bullet_match2 = re.match(r'^[-*]\s+(.*)', line)
        if bullet_match2:
            p = doc.add_paragraph(style='List Bullet')
            parts = parse_inline(bullet_match2.group(1))
            for part in parts:
                run = p.add_run(part['text'])
                run.bold = part.get('bold', False)
                run.italic = part.get('italic', False)
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraphs
        parts = parse_inline(line)
        if parts:
            p = doc.add_paragraph()
            for part in parts:
                run = p.add_run(part['text'])
                run.bold = part.get('bold', False)
                run.italic = part.get('italic', False)

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        num_cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    clean = cell_text.replace('**', '').replace('`', '').replace('*', '')
                    cell.text = clean
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    # Save
    output_path = '/home/user/ANF-OneLake-AIFoundry/misc/blog-anf-object-api-agentic-ai.docx'
    doc.save(output_path)
    print(f'Word document saved to: {output_path}')

if __name__ == '__main__':
    create_blog_docx()
